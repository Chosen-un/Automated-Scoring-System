# /rag_pipeline/retriever.py

import os
import json
import numpy as np
from typing import List, Dict, Any, Tuple, Optional 
from sentence_transformers import SentenceTransformer
import faiss
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pydantic import BaseModel, Field
from itertools import chain
import docx

from config import CHUNK_SIZE, CHUNK_OVERLAP, EMBEDDING_MODEL_NAME, FAISS_INDEX_PATH, EXTRA_DOCS_PATH
from rag_pipeline.data_processor import AssessmentRecord, FinalAssessmentRecord, QueryRewriterOutput, load_and_process_docx

class RAGDocument(BaseModel):
    id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None

class RetrievedChunk(BaseModel):
    chunk_id: str
    content: str
    metadata: Dict[str, Any]
    similarity_score: float

_embedding_model = None
_extra_docs = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        print(f"Loading embedding model: {EMBEDDING_MODEL_NAME}...")
        _embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    return _embedding_model

def get_extra_docs(extra_docs_path: str):
    global _extra_docs
    if _extra_docs is None:
        print(f"Loading extra factual documents...")
        if extra_docs_path.endswith('.json'):
            if os.path.exists(extra_docs_path):
                with open(extra_docs_path, 'r', encoding='utf-8') as f:
                    _extra_docs = json.load(f)
            else:
                print(f"Warning: JSON file not found at {extra_docs_path}.")
                _extra_docs = []
        elif extra_docs_path.endswith('.docx'):
            _extra_docs = load_and_process_docx(extra_docs_path)
        else:
            print(f"Warning: Unsupported file type for RAG documents: {extra_docs_path}")
            _extra_docs = []
        
        print(f"Loaded {len(_extra_docs)} extra factual documents.")
    return _extra_docs

def prepare_rag_corpus(rag_data: List[AssessmentRecord], extra_docs_path: str) -> Tuple[List[RAGDocument], List[RAGDocument]]:
    print("Preparing RAG corpus: chunking and embedding...")

    extra_docs_data = get_extra_docs(extra_docs_path)

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
    )
    embedding_model = get_embedding_model()

    teacher_rag_documents = []
    doc_id_counter = 0

    for record in rag_data:
        content_to_chunk = (
            f"Question: {record.question}\n"
            f"Student Answer: {record.answer}\n"
            f"Rubric: {json.dumps([c.model_dump() for c in record.rubric])}\n"
            f"Teacher Feedback: {record.feedback}\n"
            f"Awarded Score: {record.total_score}\n"
            f"Overall Quality: {record.overall_quality_of_answer}"
        )
        chunks = text_splitter.split_text(content_to_chunk)

        for i, chunk_text in enumerate(chunks):
            embedding = embedding_model.encode(chunk_text).tolist()
            rag_doc = RAGDocument(
                id=f"doc_{record.student_id}_{doc_id_counter}",
                content=chunk_text,
                metadata={
                    "student_id": record.student_id,
                    "original_question": record.question,
                    "original_answer": record.answer,
                    "chunk_index": i,
                    "total_score": record.total_score,
                    "feedback": record.feedback,
                    "rubric": [c.model_dump() for c in record.rubric],
                    "source_type": "teacher_feedback"
                },
                embedding=embedding
            )
            teacher_rag_documents.append(rag_doc)
            doc_id_counter += 1

    factual_rag_documents = []
    for doc in extra_docs_data:
        chunk_text = doc.get("text", "")
        if chunk_text.strip():
            embedding = embedding_model.encode(chunk_text).tolist()
            rag_doc = RAGDocument(
                id=f"factual_doc_{doc.get('id', 'N/A')}",
                content=chunk_text,
                metadata={
                    "source": doc.get("source", ""),
                    "source_type": "factual_doc"
                },
                embedding=embedding
            )
            factual_rag_documents.append(rag_doc)

    print(f"Teacher examples corpus prepared with {len(teacher_rag_documents)} chunks.")
    print(f"Factual documents corpus prepared with {len(factual_rag_documents)} chunks.")

    return teacher_rag_documents, factual_rag_documents

class FAISSVectorStore:
    def __init__(self, embedding_dimension: int):
        self.index = faiss.IndexFlatL2(embedding_dimension)
        self.documents: List[RAGDocument] = []

    def add_documents(self, documents: List[RAGDocument]):
        if not documents:
            return
        embeddings = np.array([doc.embedding for doc in documents]).astype('float32')
        self.index.add(embeddings)
        self.documents.extend(documents)
        print(f"Added {len(documents)} documents to FAISS index. Total: {len(self.documents)}")

    def search(self, query_embedding: List[float], k: int = 5) -> List[Tuple[RAGDocument, float]]:
        query_embedding_np = np.array(query_embedding).astype('float32').reshape(1, -1)
        distances, indices = self.index.search(query_embedding_np, k)

        results = []
        for i, idx in enumerate(indices[0]):
            if idx != -1:
                doc = self.documents[idx]
                distance = distances[0][i]
                similarity_score = -distance
                results.append((doc, similarity_score))
        return results

def perform_information_retrieval(
    record: FinalAssessmentRecord,
    teacher_faiss_store: FAISSVectorStore,
    factual_faiss_store: FAISSVectorStore,
    embedding_model: SentenceTransformer,
    top_k_teacher: int,
    top_k_factual: int
) -> FinalAssessmentRecord:
    query_text = record.query_rewriter[0].only_question if record.query_rewriter and record.query_rewriter[0].only_question else record.question
    query_embedding = embedding_model.encode(query_text).tolist()

    teacher_results = teacher_faiss_store.search(query_embedding, k=top_k_teacher)
    factual_results = factual_faiss_store.search(query_embedding, k=top_k_factual)

    all_retrieved_chunks = []
    for doc, score in chain(teacher_results, factual_results):
        all_retrieved_chunks.append(RetrievedChunk(
            chunk_id=doc.id,
            content=doc.content,
            metadata=doc.metadata,
            similarity_score=float(score)
        ))

    record.retrieved_chunks = all_retrieved_chunks
    return record